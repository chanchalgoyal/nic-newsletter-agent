import streamlit as st
import fitz
from openai import OpenAI
from langdetect import detect
from docx import Document
import tempfile

# Get API key from Streamlit secrets
client = OpenAI(api_key=st.secrets[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])

st.title("NIC Bilingual Newsletter Agent")

# -------- Agents -------- #

def reader_agent(pdf_file):
    doc = fitz.open(stream=pdf_file.read(), filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def summarize_agent(text):
    prompt = f"""
    Summarize this news in formal government newsletter tone.
    Keep it concise and factual.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": text + prompt}]
    )
    return response.choices[0].message.content

def bilingual_agent(text):
    prompt = f"""
    Convert the following into bilingual format.
    First English, then Hindi translation.
    Maintain formal government tone.
    """
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role": "user", "content": text + prompt}]
    )
    return response.choices[0].message.content

# -------- UI -------- #

uploaded_files = st.file_uploader("Upload PDF News Files", type="pdf", accept_multiple_files=True)

if st.button("Generate Newsletter") and uploaded_files:
    document = Document()
    document.add_heading("NIC Quarterly Newsletter", level=1)

    for file in uploaded_files:
        raw_text = reader_agent(file)
        summary = summarize_agent(raw_text)
        bilingual = bilingual_agent(summary)
        document.add_paragraph(bilingual)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        document.save(tmp.name)
        st.download_button("Download Newsletter", tmp.name)
